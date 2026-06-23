from __future__ import annotations

import posixpath
import shlex
import shutil
import subprocess
import tempfile
import zipfile
from dataclasses import dataclass
from html.parser import HTMLParser
from pathlib import Path
from typing import Iterable
from urllib.parse import unquote
from xml.etree import ElementTree as ET

import fitz
from charset_normalizer import from_bytes
from docx import Document as DocxDocument


@dataclass
class ExtractedText:
    chunks: list[dict]
    page_count: int = 0
    text_chars: int = 0
    has_text_layer: bool = False


def extract_pdf_text(path: Path) -> ExtractedText:
    chunks: list[dict] = []
    text_chars = 0
    with fitz.open(path) as doc:
        for page_index, page in enumerate(doc, start=1):
            text = page.get_text("text") or ""
            lines = [line.strip() for line in text.splitlines() if line.strip()]
            for line_no, line in enumerate(lines, start=1):
                chunks.append(
                    {
                        "page": page_index,
                        "ordinal": len(chunks),
                        "line": line_no,
                        "text": line,
                        "source": "pdf-text",
                    }
                )
                text_chars += len(line)
        return ExtractedText(
            chunks=chunks,
            page_count=doc.page_count,
            text_chars=text_chars,
            has_text_layer=text_chars > 0,
        )


def extract_plain_text(path: Path) -> ExtractedText:
    text = decode_text_bytes(path.read_bytes())
    return chunks_from_lines(text.splitlines(), source="text")


def extract_epub(path: Path) -> ExtractedText:
    lines: list[str] = []
    with zipfile.ZipFile(path) as archive:
        rootfile_path = _epub_rootfile_path(archive)
        opf_root = ET.fromstring(archive.read(rootfile_path))
        manifest, spine_ids = _epub_manifest_and_spine(opf_root, rootfile_path)
        html_paths = _epub_ordered_html_paths(manifest, spine_ids)

        for item_path in html_paths:
            try:
                raw = archive.read(item_path)
            except KeyError:
                continue
            lines.extend(_html_text_lines(raw))

    return chunks_from_lines(lines, source="epub")


def _epub_rootfile_path(archive: zipfile.ZipFile) -> str:
    try:
        container = ET.fromstring(archive.read("META-INF/container.xml"))
        for elem in container.iter():
            if _xml_local_name(elem.tag) == "rootfile":
                full_path = elem.attrib.get("full-path", "").strip()
                if full_path:
                    return full_path
    except (KeyError, ET.ParseError):
        pass

    for name in archive.namelist():
        if name.lower().endswith(".opf"):
            return name
    raise ValueError("EPUB package file was not found")


def _epub_manifest_and_spine(
    opf_root: ET.Element,
    rootfile_path: str,
) -> tuple[dict[str, dict[str, str]], list[str]]:
    base_dir = posixpath.dirname(rootfile_path)
    manifest: dict[str, dict[str, str]] = {}
    spine_ids: list[str] = []

    for elem in opf_root.iter():
        local_name = _xml_local_name(elem.tag)
        if local_name == "item":
            item_id = elem.attrib.get("id", "").strip()
            href = elem.attrib.get("href", "").strip()
            if not item_id or not href:
                continue
            manifest[item_id] = {
                "href": _epub_join(base_dir, href),
                "media_type": elem.attrib.get("media-type", "").strip().lower(),
            }
        elif local_name == "itemref":
            idref = elem.attrib.get("idref", "").strip()
            if idref:
                spine_ids.append(idref)
    return manifest, spine_ids


def _epub_ordered_html_paths(
    manifest: dict[str, dict[str, str]],
    spine_ids: list[str],
) -> list[str]:
    ordered: list[str] = []
    seen: set[str] = set()

    for item_id in spine_ids:
        item = manifest.get(item_id)
        if item and _epub_manifest_item_is_html(item):
            href = item["href"]
            if href not in seen:
                ordered.append(href)
                seen.add(href)

    if ordered:
        return ordered

    for item in manifest.values():
        if _epub_manifest_item_is_html(item):
            href = item["href"]
            if href not in seen:
                ordered.append(href)
                seen.add(href)
    return ordered


def _epub_manifest_item_is_html(item: dict[str, str]) -> bool:
    media_type = item.get("media_type", "")
    href = item.get("href", "").lower()
    return (
        media_type in {"application/xhtml+xml", "text/html"}
        or href.endswith(".xhtml")
        or href.endswith(".html")
        or href.endswith(".htm")
    )


def _epub_join(base_dir: str, href: str) -> str:
    path = posixpath.normpath(posixpath.join(base_dir, unquote(href.split("#", 1)[0])))
    if path.startswith("../") or path == ".." or path.startswith("/"):
        raise ValueError(f"Unsafe EPUB item path: {href}")
    return path


def _xml_local_name(tag: str) -> str:
    return tag.rsplit("}", 1)[-1]


def _html_text_lines(raw: bytes) -> list[str]:
    parser = _HtmlTextParser()
    parser.feed(decode_text_bytes(raw))
    parser.close()
    return parser.text().splitlines()


class _HtmlTextParser(HTMLParser):
    BLOCK_TAGS = {
        "address",
        "article",
        "aside",
        "blockquote",
        "body",
        "br",
        "caption",
        "dd",
        "div",
        "dl",
        "dt",
        "figcaption",
        "figure",
        "footer",
        "h1",
        "h2",
        "h3",
        "h4",
        "h5",
        "h6",
        "header",
        "hr",
        "li",
        "main",
        "nav",
        "ol",
        "p",
        "pre",
        "section",
        "table",
        "tbody",
        "td",
        "tfoot",
        "th",
        "thead",
        "tr",
        "ul",
    }
    SKIP_TAGS = {"script", "style", "noscript"}

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self._parts: list[str] = []
        self._skip_depth = 0

    def handle_starttag(self, tag: str, attrs) -> None:
        tag = tag.lower()
        if tag in self.SKIP_TAGS:
            self._skip_depth += 1
            return
        if tag in self.BLOCK_TAGS:
            self._append_break()

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        if tag in self.SKIP_TAGS:
            self._skip_depth = max(0, self._skip_depth - 1)
            return
        if tag in self.BLOCK_TAGS:
            self._append_break()

    def handle_data(self, data: str) -> None:
        if self._skip_depth or not data:
            return
        self._parts.append(data)

    def text(self) -> str:
        return "".join(self._parts)

    def _append_break(self) -> None:
        if self._parts and self._parts[-1] != "\n":
            self._parts.append("\n")


def decode_text_bytes(raw: bytes) -> str:
    if not raw:
        return ""
    candidates: list[tuple[float, str]] = []
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "big5", "big5hkscs", "cp950", "utf-16"):
        try:
            decoded = raw.decode(encoding)
        except UnicodeDecodeError:
            continue
        candidates.append((text_decode_score(decoded), decoded))

    try:
        best = from_bytes(raw).best()
        if best is not None:
            decoded = str(best)
            candidates.append((text_decode_score(decoded), decoded))
    except Exception:
        pass

    if not candidates:
        return raw.decode(errors="ignore")
    return max(candidates, key=lambda item: item[0])[1]


def text_decode_score(text: str) -> float:
    if not text:
        return 0.0
    printable = sum(1 for char in text if char.isprintable() or char in "\r\n\t")
    cjk = sum(1 for char in text if "\u3400" <= char <= "\u9fff" or "\uf900" <= char <= "\ufaff")
    replacement = text.count("\ufffd")
    controls = sum(1 for char in text if ord(char) < 32 and char not in "\r\n\t")
    return printable + cjk * 2.0 - replacement * 50.0 - controls * 20.0


def extract_docx(path: Path) -> ExtractedText:
    strict_error: Exception | None = None
    try:
        doc = DocxDocument(path)
        lines: list[str] = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                lines.append(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    lines.append(" | ".join(cells))
        return chunks_from_lines(lines, source="docx")
    except Exception as exc:
        strict_error = exc

    extracted = extract_office_text(path, source="docx")
    if extracted.has_text_layer:
        return extracted
    if strict_error is not None:
        raise strict_error
    return extracted


def extract_doc(path: Path) -> ExtractedText:
    commands = (
        ["antiword", str(path)],
        ["catdoc", str(path)],
    )
    for command in commands:
        try:
            result = subprocess.run(
                command,
                check=True,
                capture_output=True,
                text=True,
                timeout=90,
                errors="ignore",
            )
            if result.stdout.strip():
                return chunks_from_lines(result.stdout.splitlines(), source="doc")
        except (FileNotFoundError, subprocess.CalledProcessError, subprocess.TimeoutExpired):
            continue
    return extract_office_text(path, source="doc")


def extract_office_text(path: Path, source: str) -> ExtractedText:
    with tempfile.TemporaryDirectory(prefix="docsearch-office-text-") as temp:
        converted = _convert_office_document(
            path,
            Path(temp),
            convert_to="txt:Text",
            output_suffix=".txt",
            timeout=180,
        )
        if converted is None:
            return ExtractedText(chunks=[], page_count=0, text_chars=0, has_text_layer=False)
        text = decode_text_bytes(converted.read_bytes())
    return chunks_from_lines(text.splitlines(), source=source)


def chunks_from_lines(lines: Iterable[str], source: str) -> ExtractedText:
    chunks: list[dict] = []
    text_chars = 0
    for line_no, raw in enumerate(lines, start=1):
        line = " ".join(raw.strip().split())
        if not line:
            continue
        chunks.append(
            {
                "page": None,
                "ordinal": len(chunks),
                "line": line_no,
                "text": line,
                "source": source,
            }
        )
        text_chars += len(line)
    return ExtractedText(
        chunks=chunks,
        page_count=0,
        text_chars=text_chars,
        has_text_layer=text_chars > 0,
    )


def convert_office_to_pdf(path: Path, output_dir: Path) -> Path | None:
    return _convert_office_document(
        path,
        output_dir,
        convert_to="pdf",
        output_suffix=".pdf",
        timeout=180,
    )


def _convert_office_document(
    path: Path,
    output_dir: Path,
    *,
    convert_to: str,
    output_suffix: str,
    timeout: int,
) -> Path | None:
    output_dir.mkdir(parents=True, exist_ok=True)
    target = output_dir / f"{path.stem}{output_suffix}"
    target.unlink(missing_ok=True)
    source_suffix = path.suffix if path.suffix else ".bin"
    try:
        with tempfile.TemporaryDirectory(prefix="docsearch-office-source-") as temp:
            temp_dir = Path(temp)
            safe_source = temp_dir / f"source{source_suffix}"
            temp_output_dir = temp_dir / "out"
            temp_output_dir.mkdir()
            shutil.copyfile(path, safe_source)
            subprocess.run(
                [
                    "soffice",
                    "--headless",
                    "--convert-to",
                    convert_to,
                    "--outdir",
                    str(temp_output_dir),
                    str(safe_source),
                ],
                check=True,
                capture_output=True,
                text=True,
                timeout=timeout,
            )
            converted = temp_output_dir / f"{safe_source.stem}{output_suffix}"
            if not converted.exists():
                matches = list(temp_output_dir.glob(f"*{output_suffix}"))
                if not matches:
                    return None
                converted = matches[0]
            shutil.copyfile(converted, target)
            return target
    except (FileNotFoundError, subprocess.CalledProcessError, subprocess.TimeoutExpired, OSError):
        target.unlink(missing_ok=True)
        return None


class CajConversionError(RuntimeError):
    pass


def convert_caj_to_pdf(
    path: Path,
    output_dir: Path,
    command_template: str,
    timeout_seconds: int = 600,
) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    output_pdf = output_dir / f"{path.stem}.pdf"
    output_pdf.unlink(missing_ok=True)

    command = caj_command_args(command_template, path, output_pdf)
    try:
        completed = subprocess.run(
            command,
            check=True,
            capture_output=True,
            text=True,
            timeout=timeout_seconds,
            errors="replace",
        )
    except FileNotFoundError as exc:
        raise CajConversionError(
            f"CAJ converter command not found: {command[0]!r}. "
            "Install a CAJ-to-PDF converter or set CAJ_CONVERTER_COMMAND."
        ) from exc
    except subprocess.TimeoutExpired as exc:
        raise CajConversionError(f"CAJ conversion timed out after {timeout_seconds}s") from exc
    except subprocess.CalledProcessError as exc:
        detail = (exc.stderr or exc.stdout or "").strip()
        raise CajConversionError(f"CAJ conversion failed: {detail or exc}") from exc

    if not output_pdf.exists():
        detail = (completed.stderr or completed.stdout or "").strip()
        raise CajConversionError(
            f"CAJ converter did not create expected PDF: {output_pdf}. {detail}"
        )
    return output_pdf


def caj_command_args(command_template: str, input_path: Path, output_path: Path) -> list[str]:
    template = str(command_template or "").strip()
    if not template:
        raise CajConversionError("CAJ_CONVERTER_COMMAND is not configured")
    quoted_input = shlex.quote(str(input_path))
    quoted_output = shlex.quote(str(output_path))
    if "{input}" in template or "{output}" in template:
        rendered = template.format(input=quoted_input, output=quoted_output)
    else:
        rendered = f"{template} {quoted_input} {quoted_output}"
    return shlex.split(rendered)
